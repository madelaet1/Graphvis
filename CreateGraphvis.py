#********************************************************************************
#*  Script: CreateGraphvis_##.py
#*  Version: 04
#*  Author: Mike DeLaet
#*
#*  Version Notes:
#*          01 - 2-Feb 2025  | Initial development | M DeLaet
#*          02 - 7-Feb 2025  | Code cleanup and loop optimizations
#*          03 - 26-Feb 2025 | Add Diagraph    
#*          03 - 12-May 2025 | Added Class CG_internals with code documentation  
#*          04 - 15-May 2025 | Change relationship so that Nic is attached to the child of other resources
#*
#********************************************************************************

import pandas as pd
import pdb
import os
import graphviz as g
import sys
import inspect
import re
from graphviz import Digraph


#***************************************************************
# classes
#***************************************************************
class CG_internals:
        # variables
        version = "04"
        name = ""
        doc_df = pd.DataFrame(columns=["Name", "Subroutine", "Purpose"])
        document = False                # create documentation for this application
        output_file = ""                # .dot file to be created
        debug = 1                       # debugging level (0-5), 0=none, 5=verbose;
        debug_app = ""
        filtered_data = pd.DataFrame()  
        shared_xref = pd.DataFrame() 
        raw_data = pd.DataFrame()         

        # Code documentation methods
        @classmethod    
        def add(cls, name, subroutine, purpose):
            # Check if this exact row already exists
            duplicate = ((cls.doc_df["Name"] == name) &
                        (cls.doc_df["Subroutine"] == subroutine) &
                        (cls.doc_df["Purpose"] == purpose)).any()
            if not duplicate:
                cls.doc_df.loc[len(cls.doc_df)] = [name, subroutine, purpose]

                
        @classmethod
        def write_documentation_dot(cls, filename=None):
            # Generate filename from script name if not provided
            render=False
            format="svg"
            if filename is None:
                script_name = os.path.splitext(os.path.basename(sys.argv[0]))[0]
                filename = script_name

            debug_msg(1, f'Creating script documentation {filename}')

            dot = Digraph(comment="Script Documentation", format=format)

            dot.attr(label=f'{filename}')

            # First, add all unique nodes (to ensure they appear even without edges)
            for name in pd.unique(cls.doc_df["Name"]):
                dot.node(name)

            for _, row in cls.doc_df.iterrows():
                name = row["Name"]
                sub = row["Subroutine"]
                purpose = row["Purpose"]

                if sub:  # only add edges if subroutine is specified
                    dot.edge(name, sub, label=purpose)

            dot.render(filename, view=False if not render else True)

            

class LoadBalancer:
    def __init__(self, tier):
        self.tier = tier 



def cleanup_existing_files(output_dir, extensions):
    """
    Cleans up existing files in the output directory with specified extensions.
    """
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    

    for filename in os.listdir(output_dir):
        if any(filename.endswith(ext) for ext in extensions):
            os.remove(os.path.join(output_dir, filename))


def read_excel():
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    try:
        # Read both sheets from the Excel file
        CG_internals.raw_data = pd.read_excel(file_path, sheet_name=raw_data_sheet)
        resource_lookup = pd.read_excel(file_path, sheet_name=azure_services_sheet)
        CG_internals.shared_xref = pd.read_excel(file_path, sheet_name=shared_services_sheet)

        CG_internals.raw_data["AppName"] = CG_internals.raw_data["AppName"].str.strip()

        resource_lookup = resource_lookup.add_prefix("as_")
        CG_internals.shared_xref = CG_internals.shared_xref.add_prefix("xr_")

        CG_internals.shared_xref = CG_internals.shared_xref.dropna(subset=["xr_SharedAppName"])
        CG_internals.shared_xref["xr_SharedAppName"] = CG_internals.shared_xref["xr_SharedAppName"].str.strip()
        
    except Exception as e:
        error_msg(inspect.currentframe().f_code.co_name, f"Error reading Excel file: {e}")
        return

    # Merge CG_internals.raw_data with CG_internals.shared_xref on 'name' == 'SharedAppName'
    if "name" in CG_internals.raw_data.columns and "xr_PrimaryAppName" in CG_internals.shared_xref.columns:  
        CG_internals.raw_data = pd.merge(CG_internals.raw_data, CG_internals.shared_xref, left_on="AppName", right_on="xr_PrimaryAppName", how="left")
    else:
        #print(CG_internals.raw_data.columns, CG_internals.shared_xref.columns)
        debug_msg(0, "The required 'name' or 'xr_PrimaryAppName' column is missing in one of the sheets.")
        return

    # Merge RawData with ResourceLookup on the 'type' column
    if "type" in CG_internals.raw_data.columns and "as_ResourceType" in resource_lookup.columns:
        merged_data = pd.merge(CG_internals.raw_data, resource_lookup, left_on="type", right_on="as_ResourceType", how="left")
    else:
        debug_msg(0, "The required 'type' column is missing in one of the sheets.")
        return

    debug_msg(2, f"The number of rows in the merged dataset is: {merged_data.shape[0]}")

    # Filter the merged data using ResourceType and Environment
    CG_internals.filtered_data = merged_data[
        (merged_data["Environment"] == "Production")
        & (merged_data["as_Category"].str.contains("IaaS|PaaS", na=False))
    ]

def check_for_critical_columns(required_columns, data):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    
    """
    Checks if all required columns exist in the given DataFrame.
    Exits the program if any required column is missing.
    
    :param required_columns: List of critical column names.
    :param data: DataFrame to check.
    """
    for column in required_columns:
        if column not in data.columns:
            error_msg(inspect.currentframe().f_code.co_name,f"Error: Missing required column '{column}' in dataset. Exiting.")
            sys.exit(1)  # Exit with error code 1


def debug_msg(lvl, msg):
    #CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    if lvl <= CG_internals.debug:
        print(f'{msg}')

def error_msg(call_funct, msg):
    #CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    print(f'❌ {msg}.  Thrown from {call_funct}')


# Function to find the tier of a load balancer by its label
def find_loadbalancer_tier(label, df):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    row = df[df["name"] == label]
    if not row.empty:
        return row.iloc[0]["tier"]
    return None

        
def unique_name(prod_name): 
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    
    debug_msg(5, inspect.currentframe().f_code.co_name + f' 0: {prod_name}')
    left_three = prod_name[:3]
    fourth = prod_name[3]
    if "-prd-" in prod_name or "prod" in prod_name:
        ire_name = prod_name.replace("-prd-", "-ire-")
        ire_name = ire_name.replace("prod", "ire")
        dr_name = prod_name.replace("-prd-", "-drp-")
        debug_msg(5, inspect.currentframe().f_code.co_name + f" 1:{fourth}")
    elif ("com" in left_three or "crp" in left_three):        # legacy style
            ire_name = prod_name[:3] + "i" + prod_name[4:]
            dr_name = prod_name[:3] + "r" + prod_name[4:]
            debug_msg(5,"2a")
    elif left_three == "mpc":                           # greenfield 2 style
        ire_name = prod_name[:6] + "i" + prod_name[7:]
        dr_name = prod_name[:6] + "r" + prod_name[7:]
        debug_msg(5, inspect.currentframe().f_code.co_name + " 2b")  
    else:
        ire_name = prod_name + "_IRE"
        debug_msg(5, inspect.currentframe().f_code.co_name + " 3") 
        
    #print(inspect.currentframe().f_code.co_name , prod_name, ire_name)
    return ire_name

def updateUniqueNamesXls(df, uniqueXlsx_path, raw_data_sheet):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    

    try:
    # Ensure CG_internals.raw_data is a DataFrame
        if not isinstance(df, pd.DataFrame):
            raise TypeError("df must be a pandas DataFrame")

        # Ensure raw_data_sheet is defined and a string
        if not isinstance(raw_data_sheet, str) or not raw_data_sheet:
            raise ValueError("raw_data_sheet must be a non-empty string")

        #uniqueXlsx_path = "UniqueNames.xlsx"
        debug_msg(0, f'Creating {uniqueXlsx_path}\n')

        # Attempt to write to Excel
        df.to_excel(uniqueXlsx_path, sheet_name=raw_data_sheet)
        debug_msg(3, f'Successfully created {uniqueXlsx_path}')

    except PermissionError:
        error_msg(inspect.currentframe().f_code.co_name, f"Error: Cannot write to {uniqueXlsx_path}. The file might be open in another program. Close it and try again.")
    except ModuleNotFoundError as e:
        error_msg(inspect.currentframe().f_code.co_name, f"Error: {e}. Install the missing module using `pip install openpyxl` or `pip install xlsxwriter`.")
    except AttributeError as e:
        error_msg(inspect.currentframe().f_code.co_name, f"Error: {e}. Ensure raw_data is a pandas DataFrame.")
    except ValueError as e:
        error_msg(inspect.currentframe().f_code.co_name, f"Error: {e}. Sheet names must be strings with a maximum length of 31 characters.")
    except Exception as e:
        error_msg(inspect.currentframe().f_code.co_name, f"An unexpected error occurred: {e}")




def build_node_df(combined_group: pd.DataFrame, raw_data: pd.DataFrame, dot_type: str):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    # Write each resource as a node
        
    node_df = pd.DataFrame(columns=['type', 'dot_label', 'name', 'tier', 'parent', 'parent_name', 'parsed'])

    # Create node_df with attributes for the Graphviz diagrams
    for resource_type, resources in combined_group.groupby("type"):
        dot_type = resource_type.split('/')[-1]
        resource_cnt = 0
        all_cnt=0
        for idx, row in resources.iterrows():
            resource_cnt += 1
            all_cnt += 1
            label = row["name"]
            unique_flag = row.get("Unique", "N")  # Default to "N" if column missing
            node_id = f'{dot_type}_{resource_cnt}'

            if unique_flag == "Y":
                ire_name = unique_name(label)
                CG_internals.raw_data.loc[CG_internals.raw_data['name'] == label, "IRE Name"] = ire_name
            else:
                ire_name = ""

            tier=""

            '''
            debug_msg(5, "debug_1:" + node_id + " " + row['as_Category'].lower())
            if row['as_Category'].lower() == 'paas' and not dot_type == "networkinterfaces": 
                tier='paas'
                debug_msg(5, "debug_2:" + node_id + " " + row['as_Category'].lower())
            else:
                tier = None
                debug_msg(5, "debug_3:" + node_id + " " + row['as_Category'].lower())
            '''

            new_df_row = pd.DataFrame({'type': [dot_type], 'dot_label': [node_id], 'name': label, 'tier':tier, 'unique': unique_flag, 'unique_name': ire_name})
            debug_msg(5, "0." + new_df_row)
            node_df = node_df._append(new_df_row, ignore_index=True)

            debug_msg(5, "1." + str((node_df["tier"] == "paas").sum()))

            
    node_df['parent'] = ""
    node_df['parent_name'] = ""
    node_df['parsed']= ""

    #rabbit MQ servers
    rabmq_list = ["mps2517","mps2518","mps2519","mps2520","mps2521"]
    
    '''
    "microsoft.compute/virtualmachines": ('box3d', 'lightblue', 'icons\\VM-Images-l.svg', 0.35, 'bc', 'tc'),
    "microsoft.sqlvirtualmachine/sqlvirtualmachines": ('box3d', 'lightblue', 'icons\\Sql-Server.svg', 0.35, 'bc', 'tc'),
    "microsoft.network/networkinterfaces": ('component', 'lightyellow', 'icons\\Network-Interfaces-l.svg', 0.35, 'bc', 'tc'),
    "microsoft.network/loadbalancers": ('Mdiamond', 'lawngreen', 'icons\\Load-Balancers-l.svg', 0.35, 'bc', 'tc'),
    "microsoft.web/serverfarms": ('box', 'gray', 'icons\\Server-Farm.svg', 0.35, 'bc', 'tc'),
    "microsoft.storage/storageaccounts": ('folder', 'gray', 'icons\\Storage-Accounts.svg', 0.35, 'bc', 'tc'),
    "microsoft.sql/managedinstances/databases": ('folder', 'gray', 'icons\\Managed-Database.svg', 0.35, 'bc', 'tc'),
    "microsoft.sql/managedinstances": ('folder', 'gray', 'icons\\SQL-Managed-Instance.svg', 0.35, 'bc', 'tc'),
    "microsoft.keyvault/vaults": ('folder', 'gray', 'icons\\Key-Vaults.svg', 0.35, 'bc', 'tc'),
    "microsoft.appconfiguration/configurationstores": ('folder', 'gray', 'icons\\App-Configuration.svg', 0.35, 'bc', 'tc'),
    "microsoft.cache/redis": ('folder', 'gray', 'icons\\Cache-Redis.svg', 0.35, 'bc', 'tc'),
    "microsoft.servicebus/namespaces": ('folder', 'gray', 'icons\\Notification-Hub-Namespaces.svg', 0.35, 'bc', 'tc'),
    "microsoft.containerregistry/registries": ('folder', 'gray', 'icons\\Container-Registries.svg', 0.35, 'bc', 'tc'),
    "microsoft.web/sites": ('folder', 'gray', 'icons\\Function-Apps.svg', 0.35, 'bc', 'tc'),
    "microsoft.network/privateendpoints": ('folder', 'gray', 'icons\\Private-Endpoints.svg', 0.35, 'bc', 'tc')
    '''
    node_df.loc[node_df['type'] == 'site', 'tier'] = 'app'
    node_df.loc[(node_df['type'] == 'virtualmachines') & (~node_df['name'].str.contains('sql', na=False)), 'tier'] = 'web'
    node_df.loc[(node_df['type'] == 'virtualmachines') & (node_df['name'].str.contains('sql', na=False)), 'tier'] = 'sql'
    node_df.loc[(node_df['type'] == 'virtualmachines') & (node_df['name'].isin(rabmq_list)), 'tier'] = 'app'
    node_df.loc[(node_df['type'] == 'sqlvirtualmachines'), 'tier'] = 'sql'
    node_df.loc[(node_df['type'] == "loadbalancers") & (node_df['name'].str.contains('iis')), ['tier','parent']] = ['web', 'app_url']
    node_df.loc[(node_df['type'] == "loadbalancers") & (node_df['name'].str.contains('rabmq')), 'tier'] = 'app'
    node_df.loc[(node_df['type'] == "loadbalancers") & (node_df['name'].str.contains('cluster')), 'tier'] = 'sql'
    node_df.loc[(node_df['type'] == 'networkinterfaces') & (~node_df['name'].str.contains('sql', na=False)), 'tier'] = 'web'
    node_df.loc[(node_df['type'] == 'networkinterfaces') & (node_df['name'].str.contains('sql', na=False)), 'tier'] = 'sql'
    node_df.loc[(node_df['type'] == 'networkinterfaces') & (node_df['name'].isin(rabmq_list)), 'tier'] = 'app'
    node_df.loc[(node_df['type'] == 'databases') | (node_df['type'] == 'managedinstances'), 'tier'] = 'sql'
    node_df.loc[(node_df['type'] == 'serverfarms'), 'tier'] = 'app'

    node_df.loc[(node_df['type'] == 'privateendpoints') | 
                (node_df['type'] == 'registries') |
                (node_df['type'] == 'namespaces') |
                (node_df['type'] == 'configurationstores') |
                (node_df['type'] == 'vaults') |
                (node_df['type'] == 'sites') |
                (node_df['type'] == 'storageaccounts')
                , 'tier'] = 'paas'

    # Messy, but necessary.  If we don't find a tier for the LB put it in the Web 
    node_df.loc[(node_df['type'] == "loadbalancers") & (node_df['tier']==""), 'tier'] = 'web'
    lb_df = lb_df = node_df[node_df['type'] == "loadbalancers"]
    debug_msg(5, f'{inspect.currentframe().f_code.co_name}: {lb_df}\n' )



    all_cnt = len(node_df)
    return node_df, all_cnt


def process_nics(node_df: pd.DataFrame):
    CG_internals.add(inspect.currentframe().f_code.co_name, "", "")
    CG_internals.add(inspect.currentframe().f_code.co_name, "vm_details", "")

    # Iterate over NIC rows to find matches in VM rows
    nic_df = node_df[node_df['type'] == "networkinterfaces"]
    vm_df = node_df[node_df['type'].isin(["virtualmachines", "sqlvirtualmachines"])]

    # If no NIC or VM data exists, return
    if nic_df.empty or vm_df.empty: 
        return
    
    # Turn on tracing
    if(CG_internals.name==CG_internals.debug_app): pdb.set_trace()

    def update_name_if_guid(name: str) -> str:
        guid_pattern = r'\.([a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12})$'
        # Check if the name matches the GUID pattern at the end
        debug_msg(4, f'inspect.currentframe().f_code.co_name - Checking for GUID in {name}')
        if re.search(guid_pattern, name):
            # Remove the GUID part from the name
            cleaned_name = re.sub(guid_pattern, '', name)
            debug_msg(4, f'inspect.currentframe().f_code.co_name - Removing for GUID in {cleaned_name}')
            return cleaned_name
        return name  # Return original name if no GUID is found

    # Helper function for parsing NIC names
    def parse_nic_name(name: str) -> str:   
        if ".nic" in name:
            return name.split(".nic")[0]
        if "." in name:
            return name.split('.')[0]  # Fallback
        name_split = name.split('-')
        if len(name_split) >= 3:
            if name_split[0] == "nic":
                return name_split[1]
            if name_split[2] == "iac" or name_split[1] == "nic":
                return name_split[0]
        if "_z" in name:
            return name.split('_')[0][:-3]
        return name.split('-')[0]

    # Apply the update_name_if_guid function to remove GUID if present, only for rows where type == "networkinterfaces"
    node_df.loc[node_df['type'] == "networkinterfaces", 'name'] = node_df.loc[node_df['type'] == "networkinterfaces", 'name'].apply(update_name_if_guid)

    # Apply the parsing logic to NICs, only for rows where type == "networkinterfaces"
    node_df.loc[node_df['type'] == "networkinterfaces", 'parsed'] = node_df.loc[node_df['type'] == "networkinterfaces", 'name'].apply(parse_nic_name)


    # Match NICs to Virtual Machines using the parsed names
    nic_df = node_df[node_df['type'] == "networkinterfaces"]
    nic_df = nic_df.merge(vm_df[['dot_label','name', 'tier']], left_on='parsed', right_on='name', how='left', suffixes=('', '_vm'))

    # Update the tier for NICs based on the matching VM
    # node_df.loc[node_df['type'] == "networkinterfaces", 'tier'] = nic_df['tier']

    # Update the parent fields for NICs based on the matched VM
    if not nic_df.empty:
        try:
            if 'tier_vm' in nic_df.columns:
                matching_vms = nic_df[~nic_df['tier_vm'].isnull()]  # Filter rows with valid VM matches

                for _, vm_row in matching_vms.iterrows():
                    vm_dot_label = vm_row['dot_label_vm']
                    vm_name = vm_row['name_vm']
                    vm_tier = vm_row['tier_vm']
                    nic_name = vm_row['name']
                    debug_msg(5, f"{inspect.currentframe().f_code.co_name} Match found for NIC '{nic_name}', VM '{vm_name}', dot_label '{vm_dot_label}'.")

                    # Build mask for rows to update
                    mask = (node_df['type'] == 'networkinterfaces') & (node_df['parsed'] == vm_name)
                    matched_rows = node_df[mask]

                    if matched_rows.empty:
                        debug_msg(5, f"No match found in node_df for parsed NIC name '{vm_name}'. Skipping update.")
                    else:
                        debug_msg(5, f"Updating {len(matched_rows)} row(s) in node_df for NIC '{nic_name}' with parent '{vm_dot_label}'.")
                        node_df.loc[mask, ['parent', 'parent_name', 'tier']] = vm_dot_label, vm_name, vm_tier

                debug_msg(3, f"{inspect.currentframe().f_code.co_name} Final NIC processing update:\n{node_df}")

        except KeyError as e:
            error_msg(inspect.currentframe().f_code.co_name, f"Error: Missing expected column in NIC to VM match - {e}")
        except IndexError:
            error_msg(inspect.currentframe().f_code.co_name, "Error: Attempted to access an index that does not exist in NIC to VM match.")


#v04 change
def process_vms(node_df):
    CG_internals.add(inspect.currentframe().f_code.co_name, "", "")

    # Filter VMs and Load Balancers
    vm_df = node_df[(node_df['type'] == "virtualmachines") | (node_df['type'] == "sqlvirtualmachines")]
    lb_df = node_df[node_df['type'] == "loadbalancers"]

    # Turn on tracing
    #if(CG_internals.name==CG_internals.debug_app): pdb.set_trace()

    if not vm_df.empty and not lb_df.empty:
        for vm_index, vm_row in vm_df.iterrows():
            # Attempt to find a load balancer whose name contains the last 6 characters of the VM name
            matching_lb = lb_df[lb_df['name'].apply(lambda lb_name: vm_row['name'][-6:] in lb_name if pd.notna(lb_name) else False)]
            if not matching_lb.empty:
                lb = matching_lb.iloc[0]
                node_df.at[vm_index, 'parent'] = lb['dot_label']
                node_df.at[vm_index, 'parent_name'] = lb['name']
                node_df.loc[(node_df['type'] == "loadbalancers") & (node_df['name'] == lb['name']),'tier'] = node_df.at[vm_index, 'tier']



def process_paas_resources(node_df):
    
    CG_internals.add(inspect.currentframe().f_code.co_name, "", "")

    # Get all non-NIC resources where tier is 'paas'
    paas_df = node_df[(node_df['type'] != "networkinterfaces") & (node_df['tier'] == "paas")]
    nic_df = node_df[node_df['type'] == "networkinterfaces"]

    if paas_df.empty or nic_df.empty:
        return

    for nic_index, nic_row in nic_df.iterrows():
        parsed = nic_row.get('parsed')
        if pd.isna(parsed):
            continue

        # Look for a PaaS resource whose name contains this NIC's parsed value
        matching_paas = paas_df[paas_df['name'].apply(
            lambda name: parsed in name if pd.notna(name) else False
        )]

        if not matching_paas.empty:
            paas = matching_paas.iloc[0]
            node_df.at[nic_index, 'parent'] = paas['dot_label']
            node_df.at[nic_index, 'parent_name'] = paas['name']
            node_df.at[nic_index, 'tier'] = paas['tier']


def process_lbs(node_df):
    CG_internals.add(inspect.currentframe().f_code.co_name, "", "")
    # Filter Load Balancers and VMs
    lb_df = node_df[node_df['type'] == "loadbalancers"]
    #vm_df = node_df[node_df['type'].isin(["virtualmachines", "sqlvirtualmachines"]) & node_df['parent']==""]
    vm_df = node_df[node_df['type'].isin(["virtualmachines", "sqlvirtualmachines"])]

    debug_msg (3, f'{inspect.currentframe().f_code.co_name} Number of VMs with no parent: {len(vm_df)}\n{vm_df}')

    # Turn on tracing
    if(CG_internals.name==CG_internals.debug_app): pdb.set_trace()

    if not lb_df.empty and not vm_df.empty:
        for _, lb_row in lb_df.iterrows():
            lb_name = lb_row['name']
            lb_tier = lb_row['tier']
            lb_index = lb_row.name  # actual index in node_df
            debug_msg(3,f'{inspect.currentframe().f_code.co_name}: Checking {lb_name}, {lb_tier}\n')

            # Find all VMs with no parent that are in the same tier as the LB
            matching_vms = vm_df[vm_df['tier']==lb_tier]

            for _, vm_row in matching_vms.iterrows():
                vm_name = vm_row['name']
                vm_tier = vm_row['tier']

                # Set VM's parent and parent_name
                node_df.loc[node_df['name'] == vm_name, ['parent', 'parent_name']] = lb_row['dot_label'], lb_row['name']
                debug_msg(3, f'{inspect.currentframe().f_code.co_name}: Updated {vm_name} with parent: {lb_row["dot_label"]}')


                # Set Load Balancer's tier to VM's tier
                node_df.at[lb_index, 'tier'] = vm_tier
                debug_msg(3, f'{inspect.currentframe().f_code.co_name}: Load Balancer Updated with {vm_name} - {vm_tier}')
    

# depricated
def assign_lb_tier(node_df):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    #print("5.", (node_df["tier"] == "paas").sum())
    # Set NIC parents to the Web and SQL Load Balancers
    web_lb = node_df.loc[(node_df['type'] == 'loadbalancers') & (node_df['tier'] == 'web')]
    no_tier_lb = node_df.loc[(node_df['type'] == 'loadbalancers') & (node_df['tier'].isna())]
    #print(f'Load Balancers - Web Tier: {web_lb.empty}, None: {no_tier_lb.empty}\n')
    if web_lb.empty and not no_tier_lb.empty:
        #print("No Web Load Balancer found\n")
        node_df.loc[node_df['name'] == no_tier_lb.iloc[0]['name'], 'tier'] = 'web'
    elif web_lb.empty and not no_tier_lb.empty:
        node_df = node_df._append({'type':'loadbalancers', 'dot_label':'loadbalancers_00', 'name':'dummy', 'tier':'web', 'parent':'app_url'}, ignore_index=True)
        
    web_lb = node_df.loc[(node_df['type'] == 'loadbalancers') & (node_df['tier'] == 'web')]          
    if not web_lb.empty:
        node_df.loc[(node_df['type'] == 'networkinterfaces') & (node_df['tier'] == 'web'),'parent'] = web_lb.iloc[0]['dot_label']
        
    sql_lb = node_df.loc[(node_df['type'] == 'loadbalancers') & (node_df['tier'] == 'sql')]  
    # Update NIC parent for SQL Load Balancers
    if not sql_lb.empty:
        node_df.loc[(node_df['type'] == 'networkinterfaces') & (node_df['tier'] == 'sql'),'parent'] = sql_lb.iloc[0]['dot_label']

    app_lb = node_df.loc[(node_df['type'] == 'loadbalancers') & (node_df['tier'] == 'app')]  
    # Update NIC parent for Application Load Balancers
    if not app_lb.empty:
        node_df.loc[(node_df['type'] == 'networkinterfaces') & (node_df['tier'] == 'app'),'parent'] = app_lb.iloc[0]['dot_label']


def assign_tier_based_on_parent(node_df):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    #print("6.", (node_df["tier"] == "paas").sum())
    # Find rows with no parent and no tier
    no_parent_df = node_df.loc[node_df['parent'].isna()]
    paas_df = node_df.loc[node_df['tier'] == "paas"]
    
    #print("7.", (node_df["tier"] == "paas").sum())
    if not paas_df.empty and not no_parent_df.empty:
        #print(f'Looking for PaaS {len(pass_df)}, {len(no_parent_df)}\n')
        for p_index, p_row in no_parent_df.iterrows():
            name = p_row['name']
            for index, row in paas_df.iterrows():
                #print(f'Evaluate - Name: {name}, Match: {row['name']}\n')
                if p_index != index and name in row['name']:
                    node_df.loc[index, ['parent','parent_name']] = p_row['dot_label'],p_row['name']
                    #print(f'Found - Name: {name}, Match: {row['name']}\n')
    #print("8.", (node_df["tier"] == "paas").sum())
    node_df.loc[(node_df['type'].isna()),'tier'] = 'paas'
    #print("9.", (node_df["tier"] == "paas").sum())
    


def process_with_resource_lookup(file_path, raw_data_sheet, azure_services_sheet, shared_services_sheet, output_dir):
    """
    Processes the Excel file by merging RawData and ResourceLookup, filtering, and generating Graphviz files.

    Args:
        file_path (str): Path to the Excel file.
        raw_data_sheet (str): Name of the RawData sheet.
        lookup_sheet (str): Name of the ResourceLookup sheet.
        output_dir (str): Directory to save the output Graphviz files.
        """
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    
    '''
    try:
        # Read both sheets from the Excel file
        raw_data = pd.read_excel(file_path, sheet_name=raw_data_sheet)
        resource_lookup = pd.read_excel(file_path, sheet_name=azure_services_sheet)
        CG_internals.shared_xref = pd.read_excel(file_path, sheet_name=shared_services_sheet)

        raw_data["AppName"] = raw_data["AppName"].str.strip()

        resource_lookup = resource_lookup.add_prefix("as_")
        CG_internals.shared_xref = CG_internals.shared_xref.add_prefix("xr_")

        CG_internals.shared_xref = CG_internals.shared_xref.dropna(subset=["xr_SharedAppName"])
        CG_internals.shared_xref["xr_SharedAppName"] = CG_internals.shared_xref["xr_SharedAppName"].str.strip()
        
    except Exception as e:
        error_msg(inspect.currentframe().f_code.co_name, f"Error reading Excel file: {e}")
        return
    '''
    CG_internals.add(inspect.currentframe().f_code.co_name,"process_nics","Parse Nic names")
    CG_internals.add(inspect.currentframe().f_code.co_name,"process_vms","Assign Nics to VMs")
    CG_internals.add(inspect.currentframe().f_code.co_name,"process_paas_resources","Assign Nics to PaaS Services")
    CG_internals.add(inspect.currentframe().f_code.co_name,"process_lbs","Assign VMs to LoadBalancers")
    #CG_internals.add(inspect.currentframe().f_code.co_name,"assign_lb_tier","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"build_node_df","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"assign_tier_based_on_parent","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"create_graphvis_file","Create .dot files for each application")
    CG_internals.add(inspect.currentframe().f_code.co_name,"updateUniqueNamesXls","Create unique_name.xls")
    CG_internals.add(inspect.currentframe().f_code.co_name,"check_for_critical_columns","")
    
    # Check for required columns
    required_columns = ["AppName", "type", "Environment", "as_ResourceType", "as_Category", "xr_URL", "xr_SharedAppName"]
    check_for_critical_columns(required_columns, CG_internals.filtered_data)

    debug_msg(2, f"The number of rows in the filtered dataset is: {CG_internals.filtered_data.shape[0]}")
    # Sort the filtered data
    sorted_data = CG_internals.filtered_data.sort_values(by=["AppName", "as_Category", "as_ResourceType", "name"])

    # Main Loop - Look here Timur :-)
    
    # Group by AppName and create Graphviz files
    for app_name, group in sorted_data.groupby("AppName"):
        CG_internals.name = app_name
        debug_msg(1, f"\n******************\nBegin Processing Application: {app_name}")
        output_file = os.path.join(output_dir, f"{app_name}.dot")
        with open(output_file, "w") as CG_internals.output_file:
            debug_msg(2, "Note: Creating output file:" + output_file)
            dot_type = ""
            url = ""

            # Application url from column URL
            if 'xr_URL' in group.columns:
                first_url = group['xr_URL'].iloc[0]  # Attempt to fetch first URL value

                # Check if the URL is NaN or empty
                if pd.notna(first_url) and str(first_url).strip():
                    url = first_url  # Assign only if it is valid
                else:
                    url = "https://*.mgroup.net"  # Default fallback
            else:
                url = "https://*.mgroup.net"  # Default fallback if 'URL' column is missing

            
            # Initialize shared_app_name to None
            shared_app_name = None  # Default to None

            # Check if shared_rows is not empty and contains 'xr_SharedAppName'
            if not CG_internals.shared_xref.empty and "xr_SharedAppName" in CG_internals.shared_xref.columns:
                # Match rows based on AppName and xr_PrimaryAppName
                matched_rows = CG_internals.shared_xref[CG_internals.shared_xref["xr_PrimaryAppName"] == app_name]

                # If there's a match and 'xr_SharedAppName' is not NaN or empty, assign the value
                if not matched_rows.empty:
                    first_value = matched_rows["xr_SharedAppName"].iloc[0]  # Get the first valid entry
                    
                    if pd.notna(first_value) and str(first_value).strip():  # Ensure it's not NaN or empty
                        shared_app_name = first_value  # Assign the value

            #Print the result
            debug_msg(4, f"SharedAppName: {shared_app_name}")

            # Include rows where 'name' matches 'SharedAppName'
            shared_rows = sorted_data[sorted_data["AppName"] == shared_app_name]

            combined_group = pd.concat([group, shared_rows], ignore_index=True)
            combined_group = combined_group.groupby(["as_Category", "type", "name"], as_index=False).first()
      
            
            # Create node_df with attributes for the Graphviz diagrams
            node_df, nodes_cnt = build_node_df(combined_group, CG_internals.raw_data, dot_type)
            #print(f'Nodes add: {nodes_cnt}')

            lb_df = node_df[node_df['type'] == "loadbalancers"]
            if(app_name==CG_internals.debug_app): debug_msg(1,f'Step 1: {lb_df}\n')
            process_vms(node_df)
            
            lb_df = node_df[node_df['type'] == "loadbalancers"]
            if(app_name==CG_internals.debug_app): debug_msg(1,f'Step 2: {lb_df}\n')
            process_nics(node_df)  

            lb_df = node_df[node_df['type'] == "loadbalancers"]
            if(app_name==CG_internals.debug_app): debug_msg(1,f'Step 3: {lb_df}\n')
            process_paas_resources(node_df)

            lb_df = node_df[node_df['type'] == "loadbalancers"]
            if(app_name==CG_internals.debug_app): debug_msg(1,f'Step 4: {lb_df}\n')
            process_lbs(node_df)

            #assign_lb_tier(node_df)
            lb_df = node_df[node_df['type'] == "loadbalancers"]
            if(app_name==CG_internals.debug_app): debug_msg(1,f'Step 5: {lb_df}\n')
            assign_tier_based_on_parent(node_df)
                                                                                                                           
            #print("Tail", node_df[['type', 'name', 'dot_label', 'tier', 'parent','parsed']].to_string(index=False), "\n")
            

            create_graphvis_file(app_name, url, node_df, combined_group)
            # create_runbook_doc(app_name)                      Put this in the backlog
            #create_diagram_graphs(app_name , node_df)          Put this in the backlog
            if(app_name==CG_internals.debug_app): debug_msg(1, f'{node_df}\n') 
            
            debug_msg(1, f'{node_df}\n') 
    
    uniqueXlsx_path = "UniqueNames.xlsx"
    debug_msg(3,CG_internals.doc_df)   
    updateUniqueNamesXls(CG_internals.raw_data, uniqueXlsx_path,raw_data_sheet)
    CG_internals.write_documentation_dot()

#******************************************************************************
# Function: create_graphvis_file
# Purpose: Create the .dot file for a graphvis diagram
# Calls: dot_write
#******************************************************************************
def create_graphvis_file(app_name: str, url: str, node_df: pd, group: pd):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"create_cluster","Create .dot cluster for Web, App, DB and PaaS")
    CG_internals.add(inspect.currentframe().f_code.co_name,"dot_write","")
    f = CG_internals.output_file

    resource_styles = {
        "microsoft.compute/virtualmachines": ('box3d', 'lightblue', 'icons\\VM-Images-l.svg', 0.35, 'bc', 'tc'),
        "microsoft.sqlvirtualmachine/sqlvirtualmachines": ('box3d', 'lightblue', 'icons\\Sql-Server.svg', 0.35, 'bc', 'tc'),
        "microsoft.network/networkinterfaces": ('component', 'lightyellow', 'icons\\Network-Interfaces-l.svg', 0.35, 'bc', 'tc'),
        "microsoft.network/loadbalancers": ('Mdiamond', 'lawngreen', 'icons\\Load-Balancers-l.svg', 0.35, 'bc', 'tc'),
        "microsoft.web/serverfarms": ('box', 'gray', 'icons\\Server-Farm.svg', 0.35, 'bc', 'tc'),
        "microsoft.storage/storageaccounts": ('folder', 'gray', 'icons\\Storage-Accounts.svg', 0.35, 'bc', 'tc'),
        "microsoft.sql/managedinstances/databases": ('folder', 'gray', 'icons\\Managed-Database.svg', 0.35, 'bc', 'tc'),
        "microsoft.sql/managedinstances": ('folder', 'gray', 'icons\\SQL-Managed-Instance.svg', 0.35, 'bc', 'tc'),
        "microsoft.keyvault/vaults": ('folder', 'gray', 'icons\\Key-Vaults.svg', 0.35, 'bc', 'tc'),
        "microsoft.appconfiguration/configurationstores": ('folder', 'gray', 'icons\\App-Configuration.svg', 0.35, 'bc', 'tc'),
        "microsoft.cache/redis": ('folder', 'gray', 'icons\\Cache-Redis.svg', 0.35, 'bc', 'tc'),
        "microsoft.servicebus/namespaces": ('folder', 'gray', 'icons\\Notification-Hub-Namespaces.svg', 0.35, 'bc', 'tc'),
        "microsoft.containerregistry/registries": ('folder', 'gray', 'icons\\Container-Registries.svg', 0.35, 'bc', 'tc'),
        "microsoft.web/sites": ('folder', 'gray', 'icons\\Function-Apps.svg', 0.35, 'bc', 'tc'),
        "microsoft.network/privateendpoints": ('folder', 'gray', 'icons\\Private-Endpoints.svg', 0.35, 'bc', 'tc')
    }

    url_df = node_df[node_df['tier'] == "app_url"]
    web_df = node_df[node_df['tier'] == "web"]
    app_df = node_df[node_df['tier'] == "app"]
    sql_df = node_df[node_df['tier'] == "sql"]
    paas_df = node_df[node_df['tier'] == "paas"]

    tier_clusters = {
        "00": {"filtered_df": url_df, "exclude_types": ['app_url'], "tier": 'url'},
        "01": {"filtered_df": web_df, "exclude_types": ['sqlvirtualmachines','app_url'], "tier": 'web'},
        "02": {"filtered_df": app_df, "exclude_types": [''], "tier": 'app'},
        "03": {"filtered_df": sql_df, "exclude_types": ['virtualmachines'], "tier": 'sql'},
        "04": {"filtered_df": paas_df, "exclude_types": [''], "tier": 'paas'}
    }

    default_shape, default_color, default_image = 'box', 'gray', '',  
    default_margin, default_labelloc, default_imagepos = 0.00, 'mc', 'tc'

    dot_write(f'digraph "{app_name}" {{')
    dot_write(f'\t# Produced by: {os.path.basename(__file__)}')
    dot_write('\tcompound=true;\n    style=filled;')
    dot_write('\trankdir=TB;')
    dot_write('\tcolor=deepskyblue; fillcolor=lightskyblue;')

    # **Spacing & Layout Improvements**
    dot_write('\tranksep=2.0;    # Increases vertical separation')
    dot_write('\tnodesep=1.5;    # Increases horizontal spacing')
    dot_write('\toverlap=false;  # Prevents nodes from overlapping')
    dot_write('\tsplines=true;    # Enables smoother edge routing')
    dot_write('\tpackmode="clust"; # Helps in spreading clusters evenly')

    dot_write(f'\tlabel="{app_name}";\n    labelloc=t;\n')

    # **Define Node Styles**
    for resource_type, resources in group.groupby("type"):
        dot_type = resource_type.split('/')[-1]

        shape, color, image, margin, labelloc, imagepos = resource_styles.get(
            resource_type, (default_shape, default_color, default_image, default_margin, default_labelloc, default_imagepos)
        )
        #print(f'Images: {resource_type}, {image}')
        dot_write(f"\n\t/* Resource Type: {resource_type} */")

        dot_write(f'\tnode [shape={shape} style=filled fillcolor={color} '
                  f'labelloc={labelloc} margin={margin} '
                  f'imagepos={imagepos} image="{image}"];')
        type_df = node_df[node_df['type']==dot_type]

        for index, row in type_df.iterrows():  
            if dot_type == "virtualmachines" and isinstance(row['name'], str) and "sql" in row['name'].lower():
                dot_write(f'\t#"{row["dot_label"]}" [label="{row["name"]}"];')
                node_df.loc[index, ['tier','parent','parent_name']] = "","",""
            elif dot_type == "networkinterfaces":
                # Ensure 'parent' is not empty or None for network interfaces
                if pd.notna(row['parent']) and row['parent'] != "":
                    dot_write(f'\t{row["dot_label"]} [label="{row["name"]}"];')
                else:
                    dot_write(f'\t#"{row["dot_label"]}" [label="{row["name"]}"];') 
                    node_df.loc[index, ['tier','parent','parent_name']] = "","",""
            elif row['type'] == dot_type:
                dot_write(f'\t{row["dot_label"]} [label="{row["name"]}"];')  

    # **Edge Connections**
    dot_write("\t/******************************************************/")
    dot_write("\t/*    Edge Connectors between Clusters                */")
    dot_write("\t/******************************************************/")
    dot_write("\t/* app_url to web loadbalancers */")

    app_url = url if len(url) > 0 else "https:*.mgroup.net"

    dot_write(f"\t#Web site URL")
    dot_write(f'\tapp_url [label="{app_url}" shape=none margin=0.50 image="icons\\Website-Power.svg" labelloc=bc];')
    dot_write(f"\t#Web site to Web Load Balancers")

    lb_df = web_df[(web_df['type'] == "loadbalancers")]
   
    # Create Cluster connectors    
    #if not web_df.empty:
    #    if not lb_df.empty:
    #        dot_write(f"\t\tapp_url -> {lb_df.iloc[0]['dot_label']};\t\t/* {app_url} -> {lb_df.iloc[0]['name']} {lb_df.iloc[0]['tier']}*/")
    #    else:
    dot_write(f'\t\tapp_url -> Web_top;\t\t/* {app_url} */')

    if not sql_df.empty:
        dot_write(f'\tWeb_bottom -> Sql_top [ltail=cluster01, lhead=cluster03, image=""]')
    if not app_df.empty:
        dot_write(f'\tWeb_bottom -> App_top [ltail=cluster01, lhead=cluster02, image=""]')
        dot_write(f'\tApp_bottom -> Paas_top [ltail=cluster02, lhead=cluster04, image=""]')
    else:
        dot_write(f'\tWeb_bottom -> Paas_top [ltail=cluster01, lhead=cluster04, image=""]')
        

    # Loop through the dictionary and create clusters only if the tier exists in node_df
    for cluster_id, config in tier_clusters.items():
        tier = config["tier"]

        # If tier is None, we assume it's a general category (e.g., PaaS) and always process it
        if not config['filtered_df'].empty:
            #create_cluster(config['filtered_df'], cluster_id, config["node_type"], tier, f)
            create_cluster(cluster_id, config)

    dot_write("}")

    debug_msg(1, f"Generated Graphviz file: {f.name}") 


def dot_write(str: str):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    dot_file = CG_internals.output_file

    dot_file.write(f'{str}\n')

def create_diagram_graphs(app_name: str, node_df: pd):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"find_loadbalancer_tier")
    # Diagrams library
    from diagrams import Diagram
    from diagrams.azure.compute import VM, ContainerRegistries 
    from diagrams.azure.database import SQLDatabases, ManagedDatabases, SQLManagedInstances, SQLServers, CacheForRedis
    from diagrams.azure.devops import Repos
    from diagrams.azure.security import KeyVaults
    from diagrams.azure.integration import ServiceBus
    from diagrams.azure.network import LoadBalancers, NetworkInterfaces, PrivateEndpoint
    from diagrams.azure.security import SecurityCenter
    from diagrams.azure.storage import StorageAccounts
    from diagrams.azure.web import AppServices
    from diagrams.azure.integration import AppConfiguration

    #import diagrams.azure.identity as az_db
    #print(dir(az_db))

    resource_map = {
        "loadbalancers": LoadBalancers,
        "virtualmachines": VM,
        "sqlvirtualmachines": SQLServers,
        #"networkinterfaces": NetworkInterfaces,
        "storageaccounts":StorageAccounts,
        "privateendpoints": PrivateEndpoint,
        "configurationstores":AppConfiguration,
        "redis":CacheForRedis,
        "registries":ContainerRegistries,
        "vaults":KeyVaults,
        "managedinstances":SQLManagedInstances,
        "managedinstances/databases": ManagedDatabases
    }
    # Diagrams library



    # List of node types we want to extract
    node_types = [
        "loadbalancers", "virtualmachines", "sqlvirtualmachines", "networkinterfaces",
        "storageaccounts", "privateendpoints", "configurationstores", "redis",
        "registries", "vaults", "managedinstances", "managedinstances/databases"
    ]

    # Extract node lists dynamically
    node_lists = {t: node_df[node_df["type"] == t]["name"].tolist() for t in node_types}
    
    # Establishing relationships
    with Diagram("Azure Architecture", filename=f"{app_name}", outformat="svg", show=False, graph_attr={"outputorder": "edgesfirst"}):
        # Create nodes dynamically
        node_objects = {t: [resource_map[t](name) for name in node_lists[t]] for t in node_types if t in resource_map}
        #print (f'Diagram nodes: {node_objects}\n')
        loadbalancer_nodes = node_objects.get("loadbalancers", [])

        # Ensure we have at least one Load Balancer
        loadbalancers = node_lists.get("loadbalancers", [])
        lb=None;
        sqllb=None;
        
        for lb_item in loadbalancer_nodes:                 
            lb_tier = find_loadbalancer_tier(lb_item.label, node_df)
            
           # print(f'{lb_item.label} {lb_tier}\n')

            if lb_tier == "web":
                lb = lb_item 
            elif lb_tier=="sql":
                sqllb = lb_item 
            
        vm_nodes = node_objects.get("virtualmachines", [])
        sqlserver_nodes = node_objects.get("sqlvirtualmachines", [])
        managedinstance_nodes = node_objects.get("managedinstances", [])
        managedinstance_db_nodes = node_objects.get("managedinstances/databases", [])

        # Web Load Balancer to VM
        if lb:
            for vm in vm_nodes:
                lb >> vm

        # VMs to SQL Load Balancer (if exists), otherwise direct connection to SQL Server
        if sqllb:
            for vm in vm_nodes:
                vm >> sqllb
            for sqlserver in sqlserver_nodes:
                sqllb >> sqlserver
            for managedinstance in managedinstance_nodes:
                for db in managedinstance_db_nodes:
                    managedinstance >> db
        else:
            for vm in vm_nodes:
                for sqlserver in sqlserver_nodes:
                    vm >> sqlserver
                for managedinstance in managedinstance_nodes:
                    for db in managedinstance_db_nodes:
                        managedinstance >> db                

        for _, row in node_df.dropna(subset=["parent_name"]).iterrows():
            child_name = row["name"]
            parent_name = row["parent_name"]

            if parent_name in node_objects and child_name in node_objects:
                node_objects[parent_name] >> node_objects[child_name]

        # Generic resource connections from VM nodes
        #for key, nodes in node_objects.items():
        #    if key not in ["loadbalancers", "virtualmachines", "sqlvirtualmachines", "managedinstances", "managedinstances/databases"]:
        #        for vm in vm_nodes:
        #            for node in nodes:
        #                vm >> node

        # Generate debug DOT file manually
        #import subprocess
        #subprocess.run(["dot", "-Tsvg", f"{app_name}.dot", "-o", f"debug_{app_name}.svg"])


    import subprocess
    subprocess.run(["dot", "-Tsvg", f"{app_name}.dot", "-o", f"debug_{app_name}.svg"])

def create_cluster(cluster_id, config):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"dot_write","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"add_nodes","")
    
    #node_df = filtered_df
    '''
    tier_clusters = {
        "00": {"filtered_df": url_df, "exclude_types": ['app_url'], "tier": 'url'},
        "01": {"filtered_df": web_df, "exclude_types": ['sqlvirtualmachines'], "tier": 'web'},
        "02": {"filtered_df": app_df, "exclude_types": [''], "tier": 'app'},
        "03": {"filtered_df": sql_df, "exclude_types": ['virtualmachines'], "tier": 'sql'},
        "04": {"filtered_df": paas_df, "exclude_types": [''], "tier": 'paas'}
    }
    '''
    
    tier = config['tier'].capitalize()
    cluster_label = tier + " Tier"
    
    dot_write('\n\n\t/***********************************/')
    dot_write(f'\t/* {cluster_label} */')
    dot_write('\t/***********************************/')
    dot_write(f'\tsubgraph cluster{cluster_id}')
    dot_write("\t{")
    dot_write(f'\t\tlabel="{cluster_label}";')
    dot_write(f'\t\t{tier}_top [style=invisible, label="", image=""]')
    dot_write("")

    #print(f"*******Create_cluser1: {tier}, {config['exclude_types']} ")

    tier_df = config['filtered_df']
    if config['exclude_types']:
        node_types = [item for item in tier_df['type'].unique().tolist() if item not in config['exclude_types']]
        
    else:
        node_types = tier_df['type'].unique().tolist()

    debug_msg (3, f"Included types: {node_types} ")

    for node_item in node_types:
        dot_write(f'\t\t##### {node_item}')

        # Filter nodes based on type and tier
        node_type_df = tier_df[(tier_df['type'] == node_item)]

        #print(f'{tier}, {node_item}, {len(node_type_df)}\n')  # Use len(node_types) to get the length
        add_nodes(tier_df, node_type_df, node_item, tier)
        dot_write("")

    dot_write(f'\t\t{tier}_bottom [style=invisible, label="", dir=none, image=""]')
    dot_write("\t}")

def create_cluster00(df, seq, node_types, tier):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"dot_write","")

    node_df = df
    
    if tier is None:
        cluster_label = "PaaS Services"
    else:
        cluster_label = f"{tier.capitalize()} Tier"
    
    dot_write('\n\n\t/***********************************/')
    dot_write(f'\t/* {cluster_label} */')
    dot_write('\t/***********************************/')
    dot_write(f'\tsubgraph cluster{seq}')
    dot_write("\t{")
    dot_write(f'\t\tlabel="{cluster_label}";')
    
    #lb_df = node_df[(node_df['type'] == "loadbalancers") & (node_df['tier'] == tier)]  
    #nic_df = node_df[(node_df['type'] == "networkinterfaces") & (node_df['tier'] == tier)] 
    #vm_df = node_df[(node_df['type'] == "virtualmachines") & (node_df['tier'] == tier)] 
    #sql_df = node_df[(node_df['type'] == 'sqlvirtualmachines') & (node_df['tier'] == tier)] 
    
    
    #print(f'tier: {tier}, lb:{len(lb_df)}, nic:{len(nic_df)}, vm:{len(vm_df)}, sql:{len(sql_df)}\n')
    
    
    for node_item in node_types:
        dot_write(f'\n\t\t####### {node_item}')
        
        # Filter nodes based on type and tier
        node_type_df = node_df[(node_df['type'] == node_item) & (node_df['tier'] == tier)]
        
        # Add nodes of the current type
        #print(f'{tier}, {node_item}, {len(node_type_df)}\n') 
        add_nodes(node_df, node_type_df, node_item, tier)


    # Handle nodes without a parent
    na_df = node_df[(node_df['tier'] == tier) & (pd.isna(node_df['parent']))]
    if not na_df.empty:
        dot_write("\n\t\t# No parent associated with this node")
        add_nodes(node_df, na_df, "", tier)  

    '''
    # Rows without a tier, put the PaaS services in the SQL cluster
    if tier is None:
        paas_df = node_df[(node_df['tier'].isna())] 
        if not paas_df.empty:
            
            for paas_row in paas_df.itertuples():
                if not pd.isna(paas_row.parent):
                    dot_write(f'\t\t{paas_row.parent} -> {paas_row.dot_label};\t\t/* {paas_row.parent} -> {paas_row.name} */',f)
                else:
                    dot_write(f'\t\t{paas_row.dot_label};\t\t/* {paas_row.name} {paas_row.type}*/',f)
    '''
    dot_write("\t}\n")



def add_nodes(tier_df, node_type_df, node_type, tier):
    CG_internals.add(inspect.currentframe().f_code.co_name, "", "")
    CG_internals.add(inspect.currentframe().f_code.co_name, "dot_write", "")
    
    for nrow in node_type_df.itertuples():
        # 1. Connect to parent if exists, else connect to tier_top
        if pd.notna(nrow.parent) and nrow.parent != "" and nrow.parent != "app_url":
            dot_write(f'\t\t{nrow.parent} -> {nrow.dot_label};\t\t/* {nrow.parent} -> {nrow.name} */')
        else:
            dot_write(f'\t\t{tier}_top -> {nrow.dot_label} [style=invisible, dir=none];\t\t/* No parent for {nrow.name} */')

        # 2. Connect to tier_bottom if this node is not a parent of any other
        if nrow.dot_label not in tier_df['parent'].values:
            dot_write(f'\t\t{nrow.dot_label} -> {tier}_bottom [style=invisible, dir=none];\t\t/* {nrow.name} not a parent */')



from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

def add_doc_property_field(paragraph, property_name):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    """Inserts a document property Quick Part field into the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f"DOCPROPERTY {property_name}"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_runbook_doc(app):
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    doc = Document()
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = app
    core_props.subject = "Application Recovery"
    core_props.category = "Security, Runbook"
    core_props.keywords = "Runbook, Cyber Resiliency"
    core_props.company = "AHEAD"
    core_props.author = "Michael Delaet"
    core_props.created = datetime.now()
    
    # Insert Quick Parts Document Properties as Field Codes
    doc.add_paragraph("Title: ")
    add_doc_property_field(doc.add_paragraph(), "Title")

    doc.add_paragraph("Subject: ")
    add_doc_property_field(doc.add_paragraph(), "Subject")

    doc.add_paragraph("Author: ")
    add_doc_property_field(doc.add_paragraph(), "Author")

    # Contact Information
    doc.add_paragraph("\nContact:\nMichael Delaet | Principal Consultant\n\nAHEAD\n444 W Lake Street 30th Floor\nChicago, IL 60606\n312.924.4492 (office)\n800.294.5141 (fax)\nwww.AHEAD.com")

    # Legal Notice
    doc.add_paragraph("\n\nLegal notice\nThe material in this document is the proprietary property of \nAHEAD, LLC., also referred to in this document as ‘AHEAD.’ \nThis information is sensitive and is to be shared at management \ndiscretion only within AHEAD and the company to whom AHEAD \nsubmits this document.\nAll products, trademarks, and copyrights herein are \nthe property of their respective owners. \n©2025 AHEAD, LLC. All rights reserved.")

    # Save document
    filename = f"{app.replace(' ', '_')}_Runbook.docx"
    doc.save(filename)
    debug_msg(2, f"Document '{filename}' created successfully. Open it in Word and press F9 to update fields.")



if __name__ == "__main__":
    # File path to the Excel file
    CG_internals.add(inspect.currentframe().f_code.co_name,"","")
    CG_internals.add(inspect.currentframe().f_code.co_name,"cleanup_existing_files","Remove existing .dot files")
    CG_internals.add(inspect.currentframe().f_code.co_name,"read_excel","Open and read the excel into dataframes")
    CG_internals.add(inspect.currentframe().f_code.co_name,"process_with_resource_lookup","Process the excel spreadsheet")

    file_path = "AzureExport_01.xlsx"  # Update with the actual file path

    # Names of the sheets
    raw_data_sheet = "RawData_12_09"
    azure_services_sheet = "AzureServices"
    shared_services_sheet = "Shared Application XRef"

    # Output directory for the Graphviz files
    output_dir = "./graphviz_output"
    os.makedirs(output_dir, exist_ok=True)

    # Cleanup any existing Graphviz files
    cleanup_existing_files(output_dir, [".dot", ".png"])

    read_excel()

    process_with_resource_lookup(file_path, raw_data_sheet, azure_services_sheet, shared_services_sheet, output_dir)

